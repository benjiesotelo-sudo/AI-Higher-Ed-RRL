"""Build Manuscript/AI_Higher_Ed_SR_Draft_v1.docx from Manuscript/manuscript.md.

Standard journal-manuscript conventions:
  * Times New Roman 12pt body
  * Double-spaced throughout
  * 1-inch (2.54 cm) margins all sides
  * Page numbers in footer, centred
  * Section headings numbered and bold
  * Title page with working title + draft label + date + cover note
  * Markdown tables -> Word tables with header-row formatting
  * Bold/italic inline carried over
  * PRISMA Mermaid block -> labelled placeholder box
  * PLACEHOLDER paragraphs rendered italicised, in brackets
  * Section breaks between major (level-2) sections
  * Appendix A: Corpus Summary Statistics from Manuscript/prisma_data.md

Output is gitignored under Manuscript/.
"""
from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Manuscript" / "manuscript.md"
OUT_TEMPLATE = "AI_Higher_Ed_SR_Draft_{version}.docx"

WORKING_TITLE = "A Systematic Review of AI Adoption in Higher Education, 2020–2026"
TODAY = date.today().strftime("%d %B %Y")

# ----- low-level helpers ----------------------------------------------------

def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # Apply to East-Asian fallback too so the font sticks across the doc.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Times New Roman")
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)


def set_margins(doc: Document, cm: float = 2.54) -> None:
    for section in doc.sections:
        section.top_margin = Cm(cm)
        section.bottom_margin = Cm(cm)
        section.left_margin = Cm(cm)
        section.right_margin = Cm(cm)


def add_centered_page_numbers(doc: Document) -> None:
    """Insert a `PAGE` field in every section's footer, centred."""
    for section in doc.sections:
        footer = section.footer
        # Remove the default empty paragraph if any
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        # field: { PAGE }
        fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)


# ----- markdown inline parsing ----------------------------------------------

INLINE_RE = re.compile(
    r"(?P<bolditalic>\*\*\*([^\*]+)\*\*\*)"
    r"|(?P<bold>\*\*([^\*]+)\*\*)"
    r"|(?P<italic>\*([^\*]+)\*)"
    r"|(?P<code>`([^`]+)`)"
)


def add_inline_runs(paragraph, text: str, *, default_bold=False, default_italic=False) -> None:
    """Walk a markdown-flavoured string and add styled runs to a paragraph."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.bold = default_bold or None
            run.italic = default_italic or None
        if m.group("bolditalic"):
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group("bold"):
            run = paragraph.add_run(m.group(4))
            run.bold = True
        elif m.group("italic"):
            run = paragraph.add_run(m.group(6))
            run.italic = True
        elif m.group("code"):
            run = paragraph.add_run(m.group(8))
            run.font.name = "Courier New"
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = default_bold or None
        run.italic = default_italic or None


# ----- block-level rendering -----------------------------------------------

PLACEHOLDER_RE = re.compile(r"^\*?\[PLACEHOLDER\b.*?\]\*?$", re.DOTALL)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)


def add_placeholder(doc: Document, text: str) -> None:
    """Render a [PLACEHOLDER ...] block as an italicised bracketed paragraph."""
    # Strip surrounding `*` and any newlines/whitespace
    body = text.strip()
    body = body.strip("*").strip()
    # Strip surrounding [ ]
    if body.startswith("[") and body.endswith("]"):
        body = body[1:-1].strip()
    p = doc.add_paragraph()
    run = p.add_run(f"[{body}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    add_inline_runs(p, text)


def add_blockquote(doc: Document, lines: list[str]) -> None:
    text = " ".join(l.lstrip("> ").strip() for l in lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    run.italic = True


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    add_inline_runs(p, text)


def add_table_from_markdown(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = ""  # clear default empty para
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            add_inline_runs(p, cell_text, default_bold=(i == 0))
    # Add some breathing room after the table
    doc.add_paragraph()


def add_figure_placeholder(doc: Document) -> None:
    """Render the PRISMA Mermaid block as a labelled placeholder box (single-cell table)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(
        "[Figure 1: PRISMA 2020 Flow Diagram — see Manuscript/prisma_flow.md "
        "for the diagram. Will be inserted as a proper figure in the final version.]"
    )
    run.italic = True
    run.bold = True
    doc.add_paragraph()


# ----- markdown table parsing ----------------------------------------------

def parse_table_block(block: list[str]) -> list[list[str]]:
    """Parse a GitHub-flavoured markdown table block. Filter the separator row."""
    rows = []
    for line in block:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        # separator row: every cell is dashes (with optional alignment colons)
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells if c):
            continue
        rows.append(cells)
    return rows


# ----- main parsing loop ---------------------------------------------------

def parse_and_render(doc: Document, source: str) -> None:
    lines = source.splitlines()
    i = 0
    n = len(lines)
    skip_first_h1 = True  # skip the document-title H1; rendered on title page already
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # H1 (skip the doc-title, but emit any later H1s as level-1 headings)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            if skip_first_h1:
                skip_first_h1 = False
                i += 1
                continue
            add_heading(doc, stripped[2:].strip(), level=1); i += 1; continue

        # H2 — major section. Insert a page break first (except for the first one).
        if stripped.startswith("## ") and not stripped.startswith("### "):
            if any(p.text.strip() for p in doc.paragraphs):
                doc.add_page_break()
            add_heading(doc, stripped[3:].strip(), level=2); i += 1; continue

        # H3
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), level=3); i += 1; continue

        # Horizontal rule -> skip (section breaks handled via page-break on H2)
        if stripped in ("---", "***", "___"):
            i += 1; continue

        # Fenced code block
        if stripped.startswith("```"):
            fence_tag = stripped[3:].strip().lower()
            body_lines: list[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                body_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            if fence_tag == "mermaid":
                # Only the PRISMA diagram lives in a mermaid block; render as figure placeholder.
                add_figure_placeholder(doc)
            else:
                add_code_block(doc, body_lines)
            continue

        # Markdown table
        if stripped.startswith("|") and "|" in stripped[1:]:
            block: list[str] = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            rows = parse_table_block(block)
            add_table_from_markdown(doc, rows)
            continue

        # Blockquote (one or more consecutive `>` lines)
        if stripped.startswith("> "):
            block = []
            while i < n and lines[i].strip().startswith("> "):
                block.append(lines[i]); i += 1
            add_blockquote(doc, block)
            continue

        # Bullet list
        if stripped.startswith("- "):
            while i < n and lines[i].strip().startswith("- "):
                add_bullet(doc, lines[i].strip()[2:])
                i += 1
            continue

        # Blank line
        if not stripped:
            i += 1; continue

        # PLACEHOLDER paragraph (single-line or multi-line wrapped in `*[...]*`)
        if PLACEHOLDER_RE.match(stripped):
            add_placeholder(doc, stripped); i += 1; continue

        # Plain body paragraph: collect until blank line or a known block-marker
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt
                or nxt.startswith(("#", "|", "- ", "> ", "```", "---", "***", "___"))
                or PLACEHOLDER_RE.match(nxt)):
                break
            para_lines.append(lines[i])
            i += 1
        para_text = " ".join(l.strip() for l in para_lines).strip()
        if para_text:
            add_body_paragraph(doc, para_text)


# ----- title page + appendix ----------------------------------------------

def add_title_page(doc: Document) -> None:
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(WORKING_TITLE)
    r.bold = True
    r.font.size = Pt(20)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Systematic Review — Working Draft v1")
    r.italic = True
    r.font.size = Pt(14)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    r = p.add_run(TODAY); r.font.size = Pt(12)

    # Cover note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(2.0)
    p.paragraph_format.right_indent = Cm(2.0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Cover note: ")
    r.bold = True
    add_inline_runs(
        p,
        "This is a working draft. Sections marked [PLACEHOLDER] will be completed "
        "after full-text screening, CASP appraisal, and data extraction are finished. "
        "All numbers reflect the current corpus state.",
    )

    doc.add_page_break()


def add_appendix(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Appendix A: Corpus Summary Statistics", level=2)
    add_body_paragraph(
        doc,
        "Counts below are extracted directly from the canonical project database "
        "(`data/rrl.sqlite`) at the snapshot date shown on the title page. Each "
        "number is reproducible against `Manuscript/prisma_data.md`, which holds "
        "the underlying SQL queries.",
    )

    add_heading(doc, "Records identified per database", level=3)
    add_table_from_markdown(doc, [
        ["Source", "Records"],
        ["OpenAlex", "6,816"],
        ["ERIC", "16,177"],
        ["Semantic Scholar", "40,305"],
        ["**Total**", "**63,298**"],
    ])

    add_heading(doc, "Deduplication", level=3)
    add_table_from_markdown(doc, [
        ["Metric", "Count"],
        ["Total raw records", "63,298"],
        ["Unique canonical papers after dedup", "62,290"],
        ["Cross-database duplicates", "951"],
        ["Within-database duplicates", "57"],
        ["Surplus raw records collapsed", "1,008"],
    ])

    add_heading(doc, "Screening exclusions", level=3)
    add_table_from_markdown(doc, [
        ["Exclusion reason", "Count"],
        ["Not open-access (no retrievable OA PDF)", "31,364"],
        ["Not peer-reviewed (work-type / source-type signal)", "17,779"],
        ["Non-English", "4,834"],
        ["Off-topic (failed AI × HE token gate)", "7,215"],
        ["Non-empirical (review / editorial / conceptual)", "506"],
        ["K-12-only context", "35"],
        ["Wrong publication date (outside 2020–2026)", "0"],
        ["**Total excluded**", "**61,733**"],
    ])

    add_heading(doc, "Final corpus", level=3)
    add_table_from_markdown(doc, [
        ["Metric", "Count"],
        ["Records included after screening", "557"],
        ["PDFs successfully downloaded (in matrix)", "448"],
        ["PDFs failed (`oa_link_dead`)", "109"],
        ["**Quality tier — high_confidence**", "**72**"],
        ["**Quality tier — review_needed**", "**376**"],
        ["Post-ChatGPT papers (2023–2026)", "493"],
        ["Pre-ChatGPT papers (2020–2022)", "64"],
    ])

    add_body_paragraph(
        doc,
        "*Note.* The final \"studies included in review\" count after the manual "
        "full-text review of the 448 retrieved reports is still pending; the 72 / "
        "376 tier split is the current pre-full-text working count.",
    )


# ----- entry point ---------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(description="Build the manuscript docx from manuscript.md")
    parser.add_argument(
        "--version",
        default="v1",
        help="Output filename suffix (default: v1). Use v2 for the post-pivot rescrape.",
    )
    args = parser.parse_args()

    if not SRC.exists():
        print(f"ERROR: {SRC} not found")
        return 2

    doc = Document()
    set_margins(doc)
    set_default_font(doc)
    add_centered_page_numbers(doc)

    add_title_page(doc)
    parse_and_render(doc, SRC.read_text(encoding="utf-8"))
    add_appendix(doc)

    out = ROOT / "Manuscript" / OUT_TEMPLATE.format(version=args.version)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
